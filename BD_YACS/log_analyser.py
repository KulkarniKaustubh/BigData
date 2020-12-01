from docx import Document
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import os
import json
def mean_median_dump(log_dict, where, FILE):
    duration = [x1 - x2 for (x1, x2) in zip(log_dict['end_time'].values(), log_dict['arrival_time'].values())]
    mean_time = np.mean(duration)
    median_time = np.median(duration)
    
    data = json.load(FILE)
    if(where == 1):
        a = "JOB"
        data["job_completion_time"]["mean"] = mean_time
        data["job_completion_time"]["median"] = median_time
    else:
        a = "TASK"
        data["task_completion_time"]["mean"] = mean_time
        data["task_completion_time"]["median"] = median_time
    FILE.seek(0)
    json.dump(data, FILE, indent = 4)



'''def plotter(log_dict):

    imgname = log_dict['algo'][0] # Getting the name of the algorithm
    imgname += "_plot_image.png" # obtaining the appropriate name to save the plot later
    imgname = "img/" + imgname
    doc_heading = log_dict['algo'][0] + " Scheduling Algorithm Plot"
    doc_name = log_dict['algo'][0] + "_plot.docx"
    

    info = [] # list of tuples (('arrival_time', 'worker_id')) in order to sort in ascending order of 'arrival time'
    unique_worker_set = set()
    check_if_done = dict()
    end_time_list = []


    for i, j in zip(log_dict['arrival_time'].values(), log_dict['worker_id'].values()):
        info.append((i, j)) # populating the list of tuples
        unique_worker_set.add(str(j))

    for i, j in zip(log_dict['end_time'].values(), log_dict['worker_id'].values()):
        check_if_done[i] = j
        end_time_list.append(i)

    print(unique_worker_set)

    graph_dict = dict()
    graph_dict['time'] = list()
    for i in unique_worker_set:
        graph_dict[i] = list()           # dictionary to store the number of tasks assigned to each worker every time a task arrives for scheduling 

    
    info.sort(key = lambda x : x[0]) # sorting the list of tuples
    latest_arr_time = 0

    for j, i in info:

        latest_arr_time = j
        worker_set = unique_worker_set # set of all the workers
        graph_dict['time'].append(j) # keying in the arrival time of each task

        if graph_dict[str(i)] == []:
            graph_dict[str(i)] = [1] # when the first task of the job comes in, we initialize the task count of the worker that is alloted this task
        else:
            top = len(graph_dict[str(i)]) - 1
            graph_dict[str(i)].append(graph_dict[str(i)][top] + 1) # updating the task count for the worker that was alloted this task

        worker_set.remove(i) # remove the worker that was alloted this task from the list of all workers

        for k in worker_set:
            top = len(graph_dict[str(k)]) - 1
            if top == -1:
                graph_dict[str(k)].append(0)
            else:
                graph_dict[str(k)].append(graph_dict[str(k)][top]) #keeping the new task count of the other workers same as the previously updated task count
       
        for end_time in end_time_list:
            if end_time <= j:
                worker_id = check_if_done[end_time]
                top = len(graph_dict[str(worker_id)]) - 1
                graph_dict[str(worker_id)][top] = graph_dict[str(worker_id)][top] - 1
                end_time_list.remove(end_time) 



    
    # plotting the graph
    x_axis = list(range(1, len(info) + 1))
    l = np.array(x_axis) 
    plot = plt.subplot(111)
    plot.bar(l - 0.2, graph_dict[0], width=0.2, color='b', align='center')
    plot.bar(l, graph_dict[1], width=0.2, color='g', align='center')
    plot.bar(l + 0.2, graph_dict[2], width=0.2, color='r', align='center')
    plt.xlabel('Arrival time of a task')
    plt.ylabel('Number of tasks scheduled for each worker') 
    plt.legend(['worker 0', 'worker 1', 'worker 2'])
    plt.title('Number of tasks scheduled on each machine against time')
    plt.savefig(imgname)
    plt.show(block=False)
    plt.close()



    document = Document()
    document.add_heading(doc_heading,0)

    para_object = document.add_paragraph('Arrival time \t\t\t x-axis equivalents\n')
    for i,j in zip(graph_dict['time'], x_axis):
        para_object.add_run(str(i))
        para_object.add_run("\t\t")
        para_object.add_run(str(j))
        para_object.add_run("\n")
    
    para_object.add_run("\n\n")
    document.add_picture(imgname)
    document.save(doc_name)'''


def plotter(log_dict):

    imgname = log_dict['algo'][0] # Getting the name of the algorithm
    imgname += "_plot_image.png" # obtaining the appropriate name to save the plot later
    imgname = "img/" + imgname
    doc_heading = log_dict['algo'][0] + " Scheduling Algorithm Plot"
    doc_name = log_dict['algo'][0] + "_plot.docx"

    info = [] # list of tuples [('arrival_time', 'worker_id')] in order to sort in ascending order of 'arrival time'
    unique_worker_set = set()
    check_if_done = dict()
    end_time_list = []

    for i, j in zip(log_dict['arrival_time'].values(), log_dict['worker_id'].values()):
        info.append((i, j)) # populating the list of tuples
        unique_worker_set.add(j)

    for i, j in zip(log_dict['end_time'].values(), log_dict['worker_id'].values()):
        check_if_done[i] = j
        end_time_list.append(i)



    graph_dict = dict()
    graph_dict['time'] = list()
    for i in unique_worker_set:
        graph_dict[i] = list()           # dictionary to store the number of tasks assigned to each worker every time a task arrives for scheduling


    info.sort(key = lambda x : x[0]) # sorting the list of tuples
    latest_arr_time = 0



    for j, i in info:

        latest_arr_time = j
        worker_list = list(unique_worker_set) # set of all the workers
        graph_dict['time'].append(j) # keying in the arrival time of each task

        if graph_dict[i] == []:
            graph_dict[i] = [1] # when the first task of the job comes in, we initialize the task count of the worker that is alloted this task
        else:
            top = len(graph_dict[i]) - 1
            graph_dict[i].append(graph_dict[i][top] + 1) # updating the task count for the worker that was alloted this task
        
        worker_list.remove(i) # remove the worker that was alloted this task from the list of all workers
        for k in worker_list:
            top = len(graph_dict[k]) - 1
            if top == -1:
                graph_dict[k].append(0)
            else:
                graph_dict[k].append(graph_dict[k][top]) #keeping the new task count of the other workers same as the previously updated task count
        
        for end_time in end_time_list:
            if end_time <= j:
                worker_id = check_if_done[end_time]
                top = len(graph_dict[worker_id]) - 1
                graph_dict[worker_id][top] = graph_dict[worker_id][top] - 1
                end_time_list.remove(end_time)

        
    # plotting the graph
    x_axis = list(range(1, len(info) + 1))
    # l = np.array(x_axis) 
    # plot = plt.subplot(111)
    # plot.bar(l - 0.2, graph_dict[1], width=0.2, color='b', align='center')
    # plot.bar(l, graph_dict[2], width=0.2, color='g', align='center')
    # plot.bar(l + 0.2, graph_dict[3], width=0.2, color='r', align='center')

    legend_text_list = []
    for i in graph_dict.keys():
        if i != 'time':
            plt.plot(x_axis, graph_dict[i])
            legend_text_list.append('Worker ' + str(i))


    plt.xlabel('Arrival time of a task')
    plt.ylabel('Number of tasks scheduled for each worker') 
    # plt.legend(['worker 0', 'worker 1', 'worker 2'])
    plt.legend(legend_text_list)
    plt.title('Number of tasks scheduled on each machine against time')
    plt.savefig(imgname)
    plt.show(block=False)
    plt.close()

    document = Document()
    document.add_heading(doc_heading,0)

    para_object = document.add_paragraph('Arrival time \t\t\t x-axis equivalents\n')
    for i,j in zip(graph_dict['time'], x_axis):
        para_object.add_run(str(i))
        para_object.add_run("\t\t")
        para_object.add_run(str(j))
        para_object.add_run("\n")
    
    para_object.add_run("\n\n")
    document.add_picture(imgname)
    document.save(doc_name)



    
    

    


# Analysing logs/job_logs.csv
df = pd.read_csv("logs/job_log.csv")
if not df.empty: 
    # Converting the dataframe 'df' into a dictonary 'job_log_dict'
    job_log_dict = df.to_dict()
    # Create the scheduling algorith specific file to store the analysis obtained
    filename = job_log_dict['algo'][0]
    filename += "_logs_analysis.json"
    filename = "logs/" + filename
    FILE = open(filename, "w")
    data = {"job_completion_time":{}, "task_completion_time":{}}
    json.dump(data, FILE, indent = 4)
    FILE.close()
    FILE = open(filename, "r+")
    mean_median_dump(job_log_dict, 1, FILE)
    FILE.close()

# Analysing logs/task_logs.csv
df = pd.read_csv("logs/task_log.csv")
if not df.empty:
    # Converting the dataframe 'df' into a dictonary 'task_log_dict'
    task_log_dict = df.to_dict()
    FILE = open(filename, "r+")
    mean_median_dump(task_log_dict, 2, FILE)
    plotter(task_log_dict)
    FILE.close()
